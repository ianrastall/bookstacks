#!/usr/bin/env python3
"""Build reader, editable, typesetting, and corpus exports from Bookstacks TEI.

The TEI files remain canonical.  This script creates deterministic derivatives
under public/downloads without modifying the source corpus.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import textwrap
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree as ET

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
TEI_ROOT = ROOT / "tei"
OUTPUT_ROOT = Path(os.environ.get("BOOKSTACKS_EXPORT_ROOT", ROOT / "public" / "downloads")).resolve()
TMP_ROOT = ROOT / "tmp" / "exports"
FONT_PATH = ROOT / "src" / "assets" / "fonts" / "SourceSerif4.ttf"
FONT_LICENSE_PATH = ROOT / "src" / "assets" / "fonts" / "OFL.txt"
EXPORT_CSS = ROOT / "scripts" / "export" / "book.css"
XML_ID = "{http://www.w3.org/XML/1998/namespace}id"
XML_LANG = "{http://www.w3.org/XML/1998/namespace}lang"

PUBLISHED_AUTHORS = {
    "aristotle", "austen", "bronte-anne", "bronte-charlotte", "bronte-emily",
    "chesterton-g-k", "dickens", "dostoevsky", "eliot", "james", "plato",
    "shakespeare", "tolstoy", "turgenev",
}

LANGUAGES = {
    "eng": {"code": "en", "name": "English", "contents": "Contents", "note": "Note"},
    "fra": {"code": "fr", "name": "Français", "contents": "Sommaire", "note": "Note"},
    "grc": {"code": "grc", "name": "Ἑλληνική", "contents": "Περιεχόμενα", "note": "Ὑποσημείωσις"},
    "rus": {"code": "ru", "name": "Русский", "contents": "Содержание", "note": "Примечание"},
}

KIND_LABELS = {
    "en": {"act": "Act", "scene": "Scene", "chapter": "Chapter", "section": "Section", "book": "Book", "volume": "Volume", "part": "Part", "stave": "Stave", "introduction": "Introduction", "preface": "Preface", "prologue": "Prologue", "epilogue": "Epilogue"},
    "fr": {"act": "Acte", "scene": "Scène", "chapter": "Chapitre", "section": "Section", "book": "Livre", "volume": "Tome", "part": "Partie", "stave": "Strophe", "introduction": "Introduction", "preface": "Préface", "prologue": "Prologue", "epilogue": "Épilogue"},
    "grc": {"act": "Πρᾶξις", "scene": "Σκηνή", "chapter": "Κεφάλαιον", "section": "Τμῆμα", "book": "Βιβλίον", "volume": "Τόμος", "part": "Μέρος", "stave": "Στροφή", "introduction": "Εἰσαγωγή", "preface": "Προοίμιον", "prologue": "Πρόλογος", "epilogue": "Ἐπίλογος"},
    "ru": {"act": "Действие", "scene": "Сцена", "chapter": "Глава", "section": "Раздел", "book": "Книга", "volume": "Том", "part": "Часть", "stave": "Часть", "introduction": "Введение", "preface": "Предисловие", "prologue": "Пролог", "epilogue": "Эпилог"},
}

BLOCK_NAMES = {
    "div", "p", "sp", "lg", "castList", "list", "table", "figure", "floatingText",
    "opener", "closer", "postscript", "epigraph", "salute", "signed", "dateline", "trailer",
}
CHANGED_ASSETS: list[Path] = []


@dataclass
class Edition:
    source: Path
    author_slug: str
    work_slug: str
    tei_language: str
    locale: str
    language_name: str
    title: str
    author: str
    identifier: str
    license_text: str
    license_url: str
    root: ET.Element

    @property
    def stem(self) -> str:
        return self.source.stem

    @property
    def canonical_url(self) -> str:
        return f"https://bookstacks.org/{self.locale}/authors/{self.author_slug}/{self.work_slug}"

    @property
    def output_dir(self) -> Path:
        return OUTPUT_ROOT / self.author_slug / self.stem


def local_name(element: ET.Element) -> str:
    return element.tag.rsplit("}", 1)[-1]


def clean(value: str | None) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def direct_children(element: ET.Element, name: str | None = None) -> list[ET.Element]:
    children = list(element)
    return children if name is None else [child for child in children if local_name(child) == name]


def first_descendant(element: ET.Element, name: str) -> ET.Element | None:
    return next((node for node in element.iter() if local_name(node) == name), None)


def first_direct(element: ET.Element, name: str) -> ET.Element | None:
    return next((node for node in list(element) if local_name(node) == name), None)


def element_text(element: ET.Element | None) -> str:
    return clean("".join(element.itertext())) if element is not None else ""


def parse_edition(source: Path) -> Edition:
    match = re.match(r"^([^_]+)_(.+)_([a-z]{3})\.xml$", source.name)
    if not match:
        raise ValueError(f"Unexpected TEI filename: {source.name}")
    author_slug, work_slug, tei_language = match.groups()
    language = LANGUAGES.get(tei_language, {"code": tei_language, "name": tei_language})
    root = ET.parse(source).getroot()
    title_stmt = first_descendant(root, "titleStmt")
    title_node = first_direct(title_stmt, "title") if title_stmt is not None else None
    author_node = first_descendant(title_stmt, "author") if title_stmt is not None else None
    author_name_node = first_descendant(author_node, "persName") if author_node is not None else None
    license_node = first_descendant(root, "licence")
    identifier = root.get(XML_ID) or source.stem.replace("_", "-")
    return Edition(
        source=source,
        author_slug=author_slug,
        work_slug=work_slug,
        tei_language=tei_language,
        locale=str(language["code"]),
        language_name=str(language["name"]),
        title=element_text(title_node) or work_slug.replace("-", " ").title(),
        author=element_text(author_name_node) or element_text(author_node) or author_slug.replace("-", " ").title(),
        identifier=identifier,
        license_text=element_text(license_node),
        license_url=license_node.get("target", "") if license_node is not None else "",
        root=root,
    )


def markdown_escape(value: str) -> str:
    return re.sub(r"([\\`*_{}\[\]<>])", r"\\\1", value)


class MarkdownRenderer:
    def __init__(self, edition: Edition):
        self.edition = edition
        self.notes = {node.get(XML_ID): node for node in edition.root.iter() if local_name(node) == "note" and node.get(XML_ID)}
        self.referenced_note_ids = {
            target[1:]
            for node in edition.root.iter()
            if local_name(node) == "ref"
            for target in [clean(node.get("target"))]
            if target.startswith("#") and target[1:] in self.notes
        }
        self.note_labels: dict[str, str] = {}
        self.note_definitions: list[tuple[str, str]] = []
        self.generated_notes = 0

    def render(self) -> str:
        text = first_descendant(self.edition.root, "text")
        if text is None:
            raise ValueError(f"No TEI text in {self.edition.source}")
        blocks: list[str] = []
        for area in direct_children(text):
            name = local_name(area)
            if name == "front":
                blocks.append(f"# {self.area_label('front')}")
            elif name == "back":
                blocks.append(f"# {self.area_label('back')}")
            blocks.extend(self.render_container(area, 1 if name == "body" else 2))
        if self.note_definitions:
            for label, content in self.note_definitions:
                indented = content.replace("\n", "\n    ")
                blocks.append(f"[^{label}]: {indented}")
        return "\n\n".join(block for block in blocks if block.strip()).strip() + "\n"

    def area_label(self, area: str) -> str:
        labels = {
            "en": {"front": "Front matter", "back": "Back matter"},
            "fr": {"front": "Pages liminaires", "back": "Pages finales"},
            "grc": {"front": "Προοίμιον", "back": "Ἐπίλογος"},
            "ru": {"front": "Начальные материалы", "back": "Заключительные материалы"},
        }
        return labels.get(self.edition.locale, labels["en"])[area]

    def render_container(self, parent: ET.Element, depth: int) -> list[str]:
        blocks: list[str] = []
        for child in direct_children(parent):
            rendered = self.render_block(child, depth)
            if rendered:
                blocks.extend(rendered if isinstance(rendered, list) else [rendered])
        return blocks

    def render_block(self, node: ET.Element, depth: int) -> str | list[str]:
        name = local_name(node)
        if name == "head":
            return ""
        if name == "div":
            head = first_direct(node, "head")
            title = element_text(head) or self.division_label(node)
            blocks = [f"{'#' * min(6, max(1, depth))} {title}"]
            blocks.extend(self.render_container(node, depth + 1))
            return blocks
        if name == "p":
            return self.inline_content(node).strip()
        if name == "sp":
            blocks: list[str] = []
            speaker = first_direct(node, "speaker")
            if speaker is not None:
                blocks.append(f"**{element_text(speaker).rstrip('.').upper()}.**")
            for child in direct_children(node):
                if local_name(child) == "speaker":
                    continue
                rendered = self.render_block(child, depth)
                if rendered:
                    blocks.extend(rendered if isinstance(rendered, list) else [rendered])
            return blocks
        if name == "stage":
            return f"*[{self.inline_content(node).strip()}]*"
        if name == "lg":
            lines = [self.inline_content(line).strip() for line in direct_children(node, "l")]
            return "  \n".join(line for line in lines if line)
        if name == "l":
            return self.inline_content(node).strip()
        if name == "castList":
            blocks = []
            head = first_direct(node, "head")
            if head is not None:
                blocks.append(f"{'#' * min(6, depth)} {element_text(head)}")
            for item in direct_children(node, "castItem"):
                blocks.append(f"- {self.inline_content(item).strip()}")
            return blocks
        if name == "list":
            return [f"- {self.inline_content(item).strip()}" for item in direct_children(node, "item")]
        if name == "table":
            return self.render_table(node)
        if name == "figure":
            caption = first_descendant(node, "figDesc") or first_descendant(node, "head")
            return f"*{element_text(caption)}*" if caption is not None else ""
        if name in {"pb", "milestone", "fw", "metamark", "certainty", "shift"}:
            return ""
        if name in BLOCK_NAMES or any(local_name(child) in BLOCK_NAMES for child in direct_children(node)):
            return self.render_container(node, depth)
        content = self.inline_content(node).strip()
        return content

    def division_label(self, node: ET.Element) -> str:
        kind = clean(node.get("subtype") or node.get("type") or "section").replace("_", "-")
        label = KIND_LABELS.get(self.edition.locale, KIND_LABELS["en"]).get(kind, kind.replace("-", " ").title())
        number = clean(node.get("n"))
        return f"{label} {number}" if number and not number.lower().startswith("urn:") else label

    def render_table(self, node: ET.Element) -> str:
        rows = []
        for row in direct_children(node, "row"):
            cells = [self.inline_content(cell).strip().replace("|", "\\|") for cell in direct_children(row, "cell")]
            if cells:
                rows.append(cells)
        if not rows:
            return ""
        width = max(len(row) for row in rows)
        rows = [row + [""] * (width - len(row)) for row in rows]
        header = rows[0]
        body = rows[1:]
        output = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * width) + " |"]
        output.extend("| " + " | ".join(row) + " |" for row in body)
        return "\n".join(output)

    def inline_content(self, node: ET.Element) -> str:
        result = markdown_escape(node.text or "")
        for child in direct_children(node):
            result += self.inline(child)
            result += markdown_escape(child.tail or "")
        return result

    def inline(self, node: ET.Element) -> str:
        name = local_name(node)
        content = self.inline_content(node)
        if name == "note":
            note_id = node.get(XML_ID)
            return "" if note_id and note_id in self.referenced_note_ids else self.note_reference(node)
        if name == "ref":
            target = clean(node.get("target"))
            if target.startswith("#") and target[1:] in self.notes:
                return self.note_reference(self.notes[target[1:]])
            return f"[{content}]({target})" if target.startswith(("http://", "https://", "#")) else content
        if name in {"emph", "foreign", "title", "bibl"}:
            return f"*{content}*"
        if name == "hi":
            rendition = f"{node.get('rend', '')} {node.get('style', '')}".lower()
            if "bold" in rendition:
                return f"**{content}**"
            if "sup" in rendition:
                return f"<sup>{content}</sup>"
            if "sub" in rendition:
                return f"<sub>{content}</sub>"
            return f"*{content}*"
        if name in {"q", "quote"}:
            return f"“{content}”"
        if name == "stage":
            return f"*[{content}]*"
        if name == "lb":
            return "  \n"
        if name in {"pb", "milestone", "graphic", "fw", "metamark", "certainty", "shift"}:
            return ""
        if name == "gap":
            return "[…]"
        if name == "choice":
            for preferred_name in ("corr", "reg", "expan", "orig", "sic", "abbr"):
                preferred = first_direct(node, preferred_name)
                if preferred is not None:
                    return self.inline_content(preferred)
        if name == "del":
            return f"~~{content}~~"
        return content

    def note_reference(self, note: ET.Element) -> str:
        source_id = note.get(XML_ID) or ""
        if source_id and source_id in self.note_labels:
            return f"[^{self.note_labels[source_id]}]"
        self.generated_notes += 1
        label = str(self.generated_notes)
        if source_id:
            self.note_labels[source_id] = label
        content = clean("".join(note.itertext()))
        content = re.sub(r"^(?:\[(?:\*|\d+)\]|\*)\s*", "", content)
        self.note_definitions.append((label, content))
        return f"[^{label}]"


def yaml_quote(value: str) -> str:
    return json.dumps(value, ensure_ascii=False)


def markdown_document(edition: Edition, body: str) -> str:
    metadata = [
        "---",
        f"title: {yaml_quote(edition.title)}",
        f"author: {yaml_quote(edition.author)}",
        f"lang: {yaml_quote(edition.locale)}",
        f"identifier: {yaml_quote(edition.identifier)}",
        f"canonical: {yaml_quote(edition.canonical_url)}",
    ]
    if edition.license_url:
        metadata.append(f"license: {yaml_quote(edition.license_url)}")
    metadata.extend(["---", ""])
    return "\n".join(metadata) + body


def wrap_text(draw: ImageDraw.ImageDraw, value: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    words = value.split()
    if not words:
        return [""]
    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        candidate = f"{current} {word}"
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def fit_title(draw: ImageDraw.ImageDraw, value: str, font_path: Path, max_width: int, max_height: int) -> tuple[ImageFont.FreeTypeFont, str, int]:
    for size in range(148, 55, -4):
        font = ImageFont.truetype(str(font_path), size=size)
        lines = wrap_text(draw, value, font, max_width)
        spacing = int(size * 0.24)
        box = draw.multiline_textbbox((0, 0), "\n".join(lines), font=font, spacing=spacing, align="center", stroke_width=2)
        if box[2] - box[0] <= max_width and box[3] - box[1] <= max_height:
            return font, "\n".join(lines), spacing
    font = ImageFont.truetype(str(font_path), size=54)
    return font, "\n".join(wrap_text(draw, value, font, max_width)), 13


def create_cover(edition: Edition, target: Path) -> None:
    target.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGBA", (1600, 2560), "#541523")
    draw = ImageDraw.Draw(image, "RGBA")
    draw.rectangle((38, 38, 1562, 2522), outline=(0, 0, 0, 92), width=24)
    draw.rectangle((62, 62, 1538, 2498), outline=(0, 0, 0, 45), width=3)
    panel = (190, 655, 1410, 1855)
    draw.rectangle((panel[0] - 22, panel[1] - 22, panel[2] + 22, panel[3] + 22), outline=(8, 8, 8, 210), width=4)
    draw.rectangle(panel, fill=(248, 246, 243, 184), outline=(8, 8, 8, 235), width=13)
    draw.rectangle((panel[0] + 25, panel[1] + 25, panel[2] - 25, panel[3] - 25), outline=(8, 8, 8, 150), width=3)

    title_font, wrapped_title, spacing = fit_title(draw, edition.title, FONT_PATH, 1010, 650)
    title_box = draw.multiline_textbbox((0, 0), wrapped_title, font=title_font, spacing=spacing, align="center", stroke_width=2)
    title_x = 800
    title_y = 1160 - (title_box[3] - title_box[1]) // 2
    draw.multiline_text((title_x + 3, title_y + 5), wrapped_title, font=title_font, fill=(15, 15, 18, 150), anchor="ma", align="center", spacing=spacing, stroke_width=2, stroke_fill=(15, 15, 18, 120))
    draw.multiline_text((title_x, title_y), wrapped_title, font=title_font, fill=(184, 191, 198, 255), anchor="ma", align="center", spacing=spacing, stroke_width=1, stroke_fill=(72, 76, 82, 255))

    author_font = ImageFont.truetype(str(FONT_PATH), size=52)
    language_font = ImageFont.truetype(str(FONT_PATH), size=34)
    imprint_font = ImageFont.truetype(str(FONT_PATH), size=25)
    draw.text((800, 1650), edition.author, font=author_font, fill=(70, 72, 76, 255), anchor="mm")
    draw.text((800, 1742), edition.language_name, font=language_font, fill=(92, 92, 96, 230), anchor="mm")
    draw.text((800, 2415), "BOOKSTACKS", font=imprint_font, fill=(201, 190, 191, 185), anchor="mm", spacing=12)
    image.convert("RGB").save(target, format="PNG", optimize=True)


def build_reference_docx(target: Path) -> None:
    target.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Georgia"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, before, after in (("Heading 1", 18, 18, 10), ("Heading 2", 15, 14, 7), ("Heading 3", 12, 10, 5)):
        style = styles[style_name]
        style.font.name = "Georgia"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(84, 21, 35)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for style_name in ("Footnote Text", "Footnote Reference"):
        if style_name in styles:
            styles[style_name].font.name = "Georgia"
            styles[style_name].font.size = Pt(9)
    document.core_properties.author = "Bookstacks"
    document.core_properties.last_modified_by = "Bookstacks"
    document.save(target)


def insert_docx_cover(docx_path: Path, cover_path: Path, edition: Edition) -> None:
    document = Document(docx_path)
    cover_paragraph = document.add_paragraph()
    cover_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_paragraph.paragraph_format.space_before = Pt(0)
    cover_paragraph.paragraph_format.space_after = Pt(0)
    run = cover_paragraph.add_run()
    shape = run.add_picture(str(cover_path), width=Inches(5.25))
    drawing = shape._inline.docPr
    drawing.set("name", f"Cover of {edition.title}")
    drawing.set("descr", f"Oxblood and silver cover of {edition.title}")
    break_paragraph = document.add_paragraph()
    break_paragraph.add_run().add_break(WD_BREAK.PAGE)
    body = document._body._element
    body.insert(0, break_paragraph._p)
    body.insert(0, cover_paragraph._p)
    document.core_properties.title = edition.title
    document.core_properties.author = edition.author
    document.core_properties.subject = f"{edition.language_name} Bookstacks edition"
    document.core_properties.comments = f"Generated from canonical TEI source {edition.identifier}. {edition.canonical_url}"
    document.core_properties.last_modified_by = "Bookstacks"
    document.save(docx_path)


def run(command: list[str], *, cwd: Path | None = None) -> None:
    completed = subprocess.run(command, cwd=cwd, text=True, stdout=subprocess.PIPE, stderr=subprocess.STDOUT)
    if completed.returncode:
        raise RuntimeError(f"Command failed ({completed.returncode}): {' '.join(command)}\n{completed.stdout}")


def find_pandoc() -> str:
    executable = shutil.which("pandoc")
    if not executable:
        raise RuntimeError("Pandoc is required to build EPUB, HTML, DOCX, and LaTeX exports.")
    return executable


def find_tectonic() -> str:
    candidates = [
        os.environ.get("TECTONIC"),
        shutil.which("tectonic"),
        str(ROOT / ".tools" / "tectonic" / ("tectonic.exe" if os.name == "nt" else "tectonic")),
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return candidate
    raise RuntimeError("Tectonic is required to compile the LaTeX edition into PDF. Set TECTONIC or install it on PATH.")


def pandoc_metadata(edition: Edition, target: Path) -> None:
    target.write_text(json.dumps({
        "title": edition.title,
        "author": edition.author,
        "lang": edition.locale,
        "identifier": edition.identifier,
        "rights": edition.license_text,
        "source": edition.canonical_url,
    }, ensure_ascii=False, indent=2), encoding="utf-8")


def latex_escape(value: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}", "&": r"\&", "%": r"\%", "$": r"\$", "#": r"\#",
        "_": r"\_", "{": r"\{", "}": r"\}", "~": r"\textasciitilde{}", "^": r"\textasciicircum{}",
    }
    return "".join(replacements.get(character, character) for character in value)


def cover_tex(edition: Edition) -> str:
    title = latex_escape(edition.title)
    author = latex_escape(edition.author)
    language = latex_escape(edition.language_name)
    return rf"""
\begin{{titlepage}}
\thispagestyle{{empty}}
\begin{{tikzpicture}}[remember picture,overlay]
  \fill[bookoxblood] (current page.south west) rectangle (current page.north east);
  \draw[black,opacity=.38,line width=8pt] ([xshift=14pt,yshift=14pt]current page.south west) rectangle ([xshift=-14pt,yshift=-14pt]current page.north east);
  \draw[black,opacity=.20,line width=1pt] ([xshift=25pt,yshift=25pt]current page.south west) rectangle ([xshift=-25pt,yshift=-25pt]current page.north east);
  \fill[white,opacity=.72] ([xshift=.13\paperwidth,yshift=-.27\paperheight]current page.north west) rectangle ([xshift=-.13\paperwidth,yshift=.27\paperheight]current page.south east);
  \draw[black,line width=5pt] ([xshift=.13\paperwidth,yshift=-.27\paperheight]current page.north west) rectangle ([xshift=-.13\paperwidth,yshift=.27\paperheight]current page.south east);
  \draw[black,line width=1.2pt] ([xshift=.115\paperwidth,yshift=-.255\paperheight]current page.north west) rectangle ([xshift=-.115\paperwidth,yshift=.255\paperheight]current page.south east);
  \draw[black,line width=1.2pt] ([xshift=.145\paperwidth,yshift=-.285\paperheight]current page.north west) rectangle ([xshift=-.145\paperwidth,yshift=.285\paperheight]current page.south east);
  \node[align=center,text width=.66\paperwidth,text=black,opacity=.72] at ([xshift=1.4pt,yshift=.07\paperheight-1.4pt]current page.center) {{\fontsize{{28}}{{34}}\selectfont\bfseries {title}}};
  \node[align=center,text width=.66\paperwidth,text=booksilver] at ([yshift=.07\paperheight]current page.center) {{\fontsize{{28}}{{34}}\selectfont\bfseries {title}}};
  \node[align=center,text width=.64\paperwidth,text=black!72] at ([yshift=-.13\paperheight]current page.center) {{\fontsize{{14}}{{18}}\selectfont {author}\\[8pt]\fontsize{{10}}{{12}}\selectfont {language}}};
  \node[align=center,text=white,opacity=.62] at ([yshift=24pt]current page.south) {{\fontsize{{7}}{{8}}\selectfont\addfontfeatures{{LetterSpace=12}} BOOKSTACKS}};
\end{{tikzpicture}}
\null
\end{{titlepage}}
"""


def inject_latex_cover(tex_path: Path, edition: Edition) -> None:
    source = tex_path.read_text(encoding="utf-8")
    preamble = r"""
\usepackage{tikz}
\definecolor{bookoxblood}{HTML}{541523}
\definecolor{booksilver}{HTML}{B8BFC6}
\usepackage{bookmark}
\hypersetup{bookmarksopen=true,bookmarksnumbered=false,pdfdisplaydoctitle=true}
"""
    marker = "\\begin{document}"
    if marker not in source:
        raise RuntimeError(f"Pandoc LaTeX output has no document marker: {tex_path}")
    source = source.replace(marker, preamble + "\n" + marker + "\n" + cover_tex(edition), 1)
    tex_path.write_text(source, encoding="utf-8")


def zip_directory(source: Path, target: Path) -> None:
    target.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(target, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as archive:
        for item in sorted(source.rglob("*")):
            if item.is_file():
                archive.write(item, item.relative_to(source).as_posix())


def validate_epub(path: Path) -> dict[str, int]:
    with zipfile.ZipFile(path) as archive:
        names = archive.namelist()
        if not names or names[0] != "mimetype":
            raise RuntimeError(f"EPUB mimetype is not the first entry: {path}")
        if archive.getinfo("mimetype").compress_type != zipfile.ZIP_STORED:
            raise RuntimeError(f"EPUB mimetype is compressed: {path}")
        if archive.read("mimetype") != b"application/epub+zip":
            raise RuntimeError(f"Invalid EPUB mimetype: {path}")
        for required in ("META-INF/container.xml",):
            if required not in names:
                raise RuntimeError(f"EPUB is missing {required}: {path}")
        if not any(name.endswith(".opf") for name in names) or not any("nav" in name.lower() and name.endswith((".xhtml", ".html")) for name in names):
            raise RuntimeError(f"EPUB is missing package or navigation content: {path}")
        return {"entries": len(names)}


def validate_docx(path: Path, expects_notes: bool) -> dict[str, int | bool]:
    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        for required in ("word/document.xml", "word/styles.xml", "docProps/core.xml"):
            if required not in names:
                raise RuntimeError(f"DOCX is missing {required}: {path}")
        has_notes = "word/footnotes.xml" in names
        if expects_notes and not has_notes:
            raise RuntimeError(f"DOCX lost its true footnotes: {path}")
    document = Document(path)
    headings = sum(1 for paragraph in document.paragraphs if paragraph.style and paragraph.style.name.startswith("Heading"))
    if headings == 0:
        raise RuntimeError(f"DOCX contains no navigable headings: {path}")
    return {"headings": headings, "footnotes": has_notes}


def validate_pdf(path: Path) -> dict[str, int]:
    reader = PdfReader(path)
    pages = len(reader.pages)
    if pages < 2:
        raise RuntimeError(f"PDF has too few pages: {path}")
    try:
        outline_count = len(reader.outline)
    except Exception:
        outline_count = 0
    if outline_count == 0:
        raise RuntimeError(f"PDF contains no heading bookmarks: {path}")
    return {"pages": pages, "outline_items": outline_count}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def edition_fingerprint(edition: Edition) -> str:
    digest = hashlib.sha256()
    inputs = [
        edition.source,
        Path(__file__).resolve(),
        EXPORT_CSS,
        FONT_PATH,
        FONT_LICENSE_PATH,
        ROOT / "requirements-exports.txt",
    ]
    for path in inputs:
        digest.update(path.relative_to(ROOT).as_posix().encode("utf-8"))
        digest.update(path.read_bytes())
    return digest.hexdigest()


def expected_export_names(edition: Edition, formats: set[str]) -> set[str]:
    stem = edition.stem
    names = {f"{stem}.cover.png", f"{stem}.jsonl"}
    suffixes = {
        "md": "md",
        "epub": "epub",
        "html": "html.zip",
        "docx": "docx",
        "pdf": "pdf",
    }
    names.update(f"{stem}.{suffix}" for format_name, suffix in suffixes.items() if format_name in formats)
    if formats.intersection({"tex", "pdf"}):
        names.add(f"{stem}.latex.zip")
    return names


def cached_edition(
    edition: Edition,
    formats: set[str],
    fingerprint: str,
) -> tuple[dict[str, object], list[dict[str, object]]] | None:
    manifest_path = edition.output_dir / "manifest.json"
    if not manifest_path.is_file():
        return None
    try:
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None
    if manifest.get("generator_fingerprint") != fingerprint:
        return None
    if any(not (edition.output_dir / name).is_file() for name in expected_export_names(edition, formats)):
        return None
    jsonl_path = edition.output_dir / f"{edition.stem}.jsonl"
    try:
        chunks = [json.loads(line) for line in jsonl_path.read_text(encoding="utf-8").splitlines() if line]
    except (OSError, json.JSONDecodeError):
        return None
    return manifest, chunks


def text_without_notes(element: ET.Element) -> str:
    """Extract readable text while keeping inline tails and excluding note bodies."""
    parts = [element.text or ""]
    for child in list(element):
        if local_name(child) != "note":
            parts.append(text_without_notes(child))
        parts.append(child.tail or "")
    return clean("".join(parts))


def corpus_chunks(edition: Edition, max_chars: int = 8000) -> list[dict[str, object]]:
    text = first_descendant(edition.root, "text")
    if text is None:
        return []
    parent_map = {child: parent for parent in edition.root.iter() for child in list(parent)}
    divisions = [node for node in text.iter() if local_name(node) == "div" and not direct_children(node, "div")]
    if not divisions:
        divisions = [text]
    chunks: list[dict[str, object]] = []
    for division_index, division in enumerate(divisions, 1):
        heading = element_text(first_direct(division, "head"))
        if not heading:
            kind = clean(division.get("subtype") or division.get("type") or "section")
            heading = KIND_LABELS.get(edition.locale, KIND_LABELS["en"]).get(kind, kind.replace("_", " ").title())
            if clean(division.get("n")):
                heading += f" {clean(division.get('n'))}"
        hierarchy: list[str] = []
        parent = parent_map.get(division)
        while parent is not None and parent is not text:
            if local_name(parent) == "div":
                parent_heading = element_text(first_direct(parent, "head"))
                if parent_heading:
                    hierarchy.insert(0, parent_heading)
            parent = parent_map.get(parent)
        paragraphs: list[str] = []
        block_names = {"p", "l", "stage", "speaker", "castItem"}
        for block in division.iter():
            if local_name(block) not in block_names:
                continue
            parent = parent_map.get(block)
            if any(local_name(ancestor) in block_names for ancestor in iter_parents(parent, parent_map, division)):
                continue
            value = text_without_notes(block)
            if value:
                paragraphs.append(value)
        if not paragraphs:
            paragraphs = [element_text(division)]
        groups: list[list[str]] = [[]]
        group_length = 0
        for paragraph in paragraphs:
            if groups[-1] and group_length + len(paragraph) + 2 > max_chars:
                groups.append([])
                group_length = 0
            groups[-1].append(paragraph)
            group_length += len(paragraph) + 2
        annotations = [
            {"id": note.get(XML_ID) or "", "text": element_text(note)}
            for note in division.iter() if local_name(note) == "note" and element_text(note)
        ]
        division_id = division.get(XML_ID) or f"division-{division_index:04d}"
        for part_index, group in enumerate(groups, 1):
            chunks.append({
                "schema_version": "1.0",
                "work_id": f"{edition.author_slug}:{edition.work_slug}",
                "edition_id": edition.identifier,
                "language": edition.locale,
                "chunk_id": f"{division_id}.{part_index}",
                "parent_id": division_id,
                "kind": clean(division.get("subtype") or division.get("type") or "section"),
                "heading": heading,
                "hierarchy": hierarchy,
                "sequence": len(chunks) + 1,
                "text": "\n\n".join(group),
                "annotations": annotations,
                "source_url": edition.canonical_url,
                "license": edition.license_url or edition.license_text,
            })
    return chunks


def iter_parents(
    node: ET.Element | None,
    parent_map: dict[ET.Element, ET.Element],
    stop: ET.Element,
) -> Iterable[ET.Element]:
    while node is not None and node is not stop:
        yield node
        node = parent_map.get(node)


def write_jsonl(records: Iterable[dict[str, object]], target: Path) -> None:
    target.parent.mkdir(parents=True, exist_ok=True)
    with target.open("w", encoding="utf-8", newline="\n") as stream:
        for record in records:
            stream.write(json.dumps(record, ensure_ascii=False, separators=(",", ":")) + "\n")


def safe_reset_directory(target: Path) -> None:
    resolved = target.resolve()
    root = TMP_ROOT.resolve()
    if not resolved.is_relative_to(root) or resolved == root:
        raise RuntimeError(f"Refusing to reset directory outside task temp root: {resolved}")
    if resolved.exists():
        shutil.rmtree(resolved)
    resolved.mkdir(parents=True, exist_ok=True)


def build_edition(edition: Edition, formats: set[str]) -> tuple[dict[str, object], list[dict[str, object]]]:
    edition.output_dir.mkdir(parents=True, exist_ok=True)
    fingerprint = edition_fingerprint(edition)
    cached = cached_edition(edition, formats, fingerprint)
    if cached is not None:
        print("  reused validated exports", flush=True)
        return cached
    work_dir = TMP_ROOT / edition.stem
    safe_reset_directory(work_dir)
    body = MarkdownRenderer(edition).render()
    content_md = work_dir / "content.md"
    content_md.write_text(body, encoding="utf-8", newline="\n")
    metadata_path = work_dir / "metadata.json"
    pandoc_metadata(edition, metadata_path)
    cover_path = edition.output_dir / f"{edition.stem}.cover.png"
    create_cover(edition, cover_path)
    chunks = corpus_chunks(edition)
    edition_jsonl = edition.output_dir / f"{edition.stem}.jsonl"
    write_jsonl(chunks, edition_jsonl)
    generated: list[Path] = [cover_path, edition_jsonl]
    validations: dict[str, object] = {"jsonl_chunks": len(chunks)}
    pandoc = find_pandoc() if formats.intersection({"epub", "html", "docx", "tex", "pdf"}) else ""

    if "md" in formats:
        markdown_path = edition.output_dir / f"{edition.stem}.md"
        markdown_path.write_text(markdown_document(edition, body), encoding="utf-8", newline="\n")
        generated.append(markdown_path)

    if "epub" in formats:
        epub_path = edition.output_dir / f"{edition.stem}.epub"
        run([pandoc, str(content_md), "--from=markdown+smart", "--to=epub3", f"--output={epub_path}", "--toc", f"--epub-cover-image={cover_path}", f"--css={EXPORT_CSS}", f"--metadata-file={metadata_path}"])
        validations["epub"] = validate_epub(epub_path)
        generated.append(epub_path)

    if "html" in formats:
        html_dir = work_dir / "html"
        html_dir.mkdir()
        html_input = work_dir / "html-content.md"
        html_input.write_text(f"![{markdown_escape(edition.title)}]({cover_path.as_posix()}){{.book-cover}}\n\n{body}", encoding="utf-8")
        html_path = html_dir / "index.html"
        run([pandoc, str(html_input), "--from=markdown+smart", "--to=html5", f"--output={html_path}", "--standalone", "--toc", "--embed-resources", f"--css={EXPORT_CSS}", f"--metadata-file={metadata_path}"])
        (html_dir / "README.txt").write_text(f"Offline HTML edition of {edition.title}.\nCanonical edition: {edition.canonical_url}\n", encoding="utf-8")
        html_zip = edition.output_dir / f"{edition.stem}.html.zip"
        zip_directory(html_dir, html_zip)
        generated.append(html_zip)

    if "docx" in formats:
        reference_docx = work_dir / "bookstacks-reference.docx"
        build_reference_docx(reference_docx)
        docx_path = edition.output_dir / f"{edition.stem}.docx"
        run([pandoc, str(content_md), "--from=markdown+smart", "--to=docx", f"--output={docx_path}", "--toc", f"--reference-doc={reference_docx}", f"--metadata-file={metadata_path}"])
        insert_docx_cover(docx_path, cover_path, edition)
        validations["docx"] = validate_docx(docx_path, "[^" in body)
        generated.append(docx_path)

    if formats.intersection({"tex", "pdf"}):
        latex_dir = work_dir / "latex"
        latex_dir.mkdir()
        shutil.copy2(FONT_PATH, latex_dir / FONT_PATH.name)
        shutil.copy2(FONT_LICENSE_PATH, latex_dir / FONT_LICENSE_PATH.name)
        shutil.copy2(cover_path, latex_dir / cover_path.name)
        tex_path = latex_dir / f"{edition.stem}.tex"
        run([
            pandoc, str(content_md), "--from=markdown+smart", "--to=latex", "--standalone", f"--output={tex_path}", "--toc",
            "--top-level-division=chapter", "--variable=documentclass:book", "--variable=papersize:letter",
            "--variable=classoption:oneside", "--variable=classoption:openany",
            "--variable=fontsize:11pt", f"--variable=mainfont:{FONT_PATH.name}", f"--variable=title-meta:{edition.title}",
            f"--variable=author-meta:{edition.author}", f"--variable=lang:{edition.locale}", "--variable=colorlinks:true",
        ])
        inject_latex_cover(tex_path, edition)
        (latex_dir / "README.md").write_text(
            f"# {edition.title}\n\nGenerated from `{edition.source.name}`. The TEI source remains canonical.\n\n"
            f"Compile with `tectonic {tex_path.name}`. The bundled Source Serif 4 font is licensed under the SIL Open Font License.\n\n"
            f"Canonical edition: {edition.canonical_url}\n",
            encoding="utf-8",
        )
        if "pdf" in formats:
            tectonic = find_tectonic()
            run([tectonic, "--keep-logs", "--outdir", str(latex_dir), tex_path.name], cwd=latex_dir)
            compiled_pdf = latex_dir / f"{edition.stem}.pdf"
            pdf_path = edition.output_dir / compiled_pdf.name
            shutil.copy2(compiled_pdf, pdf_path)
            validations["pdf"] = validate_pdf(pdf_path)
            generated.append(pdf_path)
        latex_zip = edition.output_dir / f"{edition.stem}.latex.zip"
        zip_directory(latex_dir, latex_zip)
        generated.append(latex_zip)

    manifest = {
        "schema_version": "1.0",
        "edition_id": edition.identifier,
        "work_id": f"{edition.author_slug}:{edition.work_slug}",
        "title": edition.title,
        "author": edition.author,
        "language": edition.locale,
        "canonical_url": edition.canonical_url,
        "source_tei": edition.source.name,
        "generator_fingerprint": fingerprint,
        "validations": validations,
        "files": [
            {"name": path.name, "bytes": path.stat().st_size, "sha256": sha256(path)}
            for path in sorted(generated)
        ],
    }
    manifest_path = edition.output_dir / "manifest.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    CHANGED_ASSETS.extend((*generated, manifest_path))
    return manifest, chunks


def source_files(arguments: argparse.Namespace) -> list[Path]:
    if arguments.file:
        files = [(ROOT / value).resolve() if not Path(value).is_absolute() else Path(value).resolve() for value in arguments.file]
    else:
        files = sorted(path for path in TEI_ROOT.glob("*/*.xml") if path.parent.name in PUBLISHED_AUTHORS)
    for path in files:
        if not path.is_file() or path.suffix.lower() != ".xml" or not path.is_relative_to(TEI_ROOT.resolve()):
            raise ValueError(f"Export input must be a TEI XML file inside {TEI_ROOT}: {path}")
    return files


def write_collection_exports(manifests: list[dict[str, object]], chunks: list[dict[str, object]]) -> None:
    corpus_dir = OUTPUT_ROOT / "corpus"
    corpus_dir.mkdir(parents=True, exist_ok=True)
    jsonl_path = corpus_dir / "bookstacks-corpus.jsonl"
    write_jsonl(chunks, jsonl_path)
    manifest = {
        "schema_version": "1.0",
        "edition_count": len(manifests),
        "chunk_count": len(chunks),
        "languages": sorted({str(item["language"]) for item in manifests}),
        "editions": manifests,
        "files": [{"name": jsonl_path.name, "bytes": jsonl_path.stat().st_size, "sha256": sha256(jsonl_path)}],
    }
    manifest_path = corpus_dir / "bookstacks-corpus-manifest.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    schema_path = corpus_dir / "README.md"
    schema_path.write_text(
        "# Bookstacks corpus export\n\nEach JSONL record is a paragraph-boundary chunk within a semantic TEI division. "
        "Stable work, edition, division, language, hierarchy, annotation, source, and license fields are retained. "
        "The TEI editions remain canonical.\n",
        encoding="utf-8",
    )
    archive_path = corpus_dir / "bookstacks-corpus.zip"
    with zipfile.ZipFile(archive_path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as archive:
        for path in (jsonl_path, manifest_path, schema_path):
            archive.write(path, path.name)
    CHANGED_ASSETS.extend((jsonl_path, manifest_path, schema_path, archive_path))


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--file", action="append", help="Build one TEI file (relative to the repository or absolute).")
    parser.add_argument("--all", action="store_true", help="Build every published TEI edition (the default when --file is omitted).")
    parser.add_argument("--formats", default="md,epub,html,docx,tex,pdf", help="Comma-separated formats: md,epub,html,docx,tex,pdf")
    arguments = parser.parse_args()
    formats = {value.strip() for value in arguments.formats.split(",") if value.strip()}
    allowed = {"md", "epub", "html", "docx", "tex", "pdf"}
    unknown = formats - allowed
    if unknown:
        parser.error(f"Unknown formats: {', '.join(sorted(unknown))}")
    if "pdf" in formats:
        formats.add("tex")
    TMP_ROOT.mkdir(parents=True, exist_ok=True)
    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)
    CHANGED_ASSETS.clear()
    files = source_files(arguments)
    manifests: list[dict[str, object]] = []
    all_chunks: list[dict[str, object]] = []
    for index, source in enumerate(files, 1):
        edition = parse_edition(source)
        print(f"[{index}/{len(files)}] {edition.title} ({edition.locale})", flush=True)
        manifest, chunks = build_edition(edition, formats)
        manifests.append(manifest)
        all_chunks.extend(chunks)
    if not arguments.file:
        write_collection_exports(manifests, all_chunks)
    changed_path = OUTPUT_ROOT / "_changed-assets.json"
    changed_path.write_text(
        json.dumps([path.relative_to(OUTPUT_ROOT).as_posix() for path in CHANGED_ASSETS], indent=2) + "\n",
        encoding="utf-8",
    )
    print(f"Built {len(manifests)} edition export set(s).", flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
